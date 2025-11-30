"""Export utilities for OCR results."""
from __future__ import annotations

from pathlib import Path
from typing import List

import docx


def export_txt(output_path: Path, pages: List[str]) -> None:
    """Save pages to a TXT file."""
    output_path.write_text("\n\n".join(pages), encoding="utf-8")


def export_docx(output_path: Path, pages: List[str]) -> None:
    """Save pages to a DOCX file."""
    document = docx.Document()
    for idx, page in enumerate(pages):
        document.add_heading(f"Page {idx + 1}", level=2)
        document.add_paragraph(page)
    document.save(output_path)
